import os, sys
import torch
from transformers import GPT2LMHeadModel, GPT2Tokenizer, AdamW, get_linear_schedule_with_warmup
from googlesearch import search
import requests
from bs4 import BeautifulSoup
import PyPDF2
import docx
import pandas as pd
import tempfile
import shutil

# ==================== تابع آموزش مدل (Create_Model) ====================
def Create_Model(robot_name, folder_path, accuracy_param=0.9, test_size_param=0.1, number_test=3):
    try:
        os.chdir("Models")
    except:
        pass
    texts = []
    # برای هر فایل در پوشه اطلاعات
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        # استخراج متن از PDF
        if filename.lower().endswith('.pdf'):
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
                texts.append(text)
        # استخراج متن از فایل Word
        elif filename.lower().endswith('.docx'):
            doc = docx.Document(file_path)
            text = ""
            for para in doc.paragraphs:
                text += para.text
            texts.append(text)
        # استخراج داده‌ها از Excel (xlsx)
        elif filename.lower().endswith('.xlsx'):
            df = pd.read_excel(file_path)
            text = df.to_string()  # تبدیل داده‌ها به رشته
            texts.append(text)
        # استخراج متن از فایل TXT و CSV
        elif filename.lower().endswith('.txt') or filename.lower().endswith('.csv'):
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            texts.append(text)
        # استخراج متن از فایل صوتی (در اینجا از تابع analyze_audio استفاده می‌شود)
        elif filename.lower().endswith(('.mp3', '.wav', '.aac', '.ogg')):
            with open(file_path, 'rb') as f:
                text = analyze_audio(f)
            texts.append(text)
    # ترکیب تمامی متن‌ها
    data = " ".join(texts)
    # پیش‌پردازش داده‌ها و آموزش مدل
    tokenizer = GPT2Tokenizer.from_pretrained('gpt2')
    tokenizer.pad_token = tokenizer.eos_token
    model = GPT2LMHeadModel.from_pretrained('gpt2')
    inputs = tokenizer(data, return_tensors="pt", max_length=1024, truncation=True, padding=True)
    if len(inputs['input_ids'][0]) < 2:
        raise ValueError("تعداد داده‌ها کافی نیست. لطفاً داده‌های بیشتری وارد کنید.")
    from sklearn.model_selection import train_test_split
    train_data, val_data = train_test_split(inputs['input_ids'].squeeze().tolist(), test_size=test_size_param)
    from torch.utils.data import DataLoader, Dataset
    class TextDataset(Dataset):
        def __init__(self, input_ids):
            self.input_ids = input_ids
        def __len__(self):
            return len(self.input_ids)
        def __getitem__(self, idx):
            return {'input_ids': torch.tensor(self.input_ids[idx])}
    train_dataset = TextDataset(train_data)
    val_dataset = TextDataset(val_data)
    train_dataloader = DataLoader(train_dataset, batch_size=4, shuffle=True)
    val_dataloader = DataLoader(val_dataset, batch_size=4, shuffle=False)
    optimizer = AdamW(model.parameters(), lr=5e-5)
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.to(device)
    num_train_steps = len(train_dataloader) * 3
    lr_scheduler = get_linear_schedule_with_warmup(optimizer, num_warmup_steps=0, num_training_steps=num_train_steps)
    print("Start Learn Model...")
    model.train()
    for epoch in range(number_test):
        for batch in train_dataloader:
            optimizer.zero_grad()
            batch = {key: value.to(device) for key, value in batch.items()}
            outputs = model(**batch, labels=batch['input_ids'])
            loss = outputs.loss
            loss.backward()
            optimizer.step()
            lr_scheduler.step()
            print(f"Epoch {epoch+1}, Loss: {loss.item():.4f}")
        # ذخیره مدل پس از هر اپوک
        model.save_pretrained(f"./{robot_name}_epoch_{epoch+1}")
        tokenizer.save_pretrained(f"./{robot_name}_epoch_{epoch+1}")
    output_folder = os.path.join(os.getcwd(), robot_name)
    os.makedirs(output_folder, exist_ok=True)
    model.save_pretrained(output_folder)
    tokenizer.save_pretrained(output_folder)
    print(f"{robot_name} Model Saved In {output_folder}")

# ==================== توابع تحلیل فایل ====================
def analyze_image(file_obj):
    try:
        from PIL import Image
        import pytesseract
    except ImportError:
        return "امکان تحلیل تصویر فراهم نیست (کتابخانه‌های مورد نیاز نصب نشده‌اند)."
    try:
        image = Image.open(file_obj)
        text = pytesseract.image_to_string(image, lang='fas')
        description = "این تصویر شامل یک صحنه زیبا با رنگ‌های دلنشین است."
        if text.strip():
            return f"تحلیل تصویر: {description} متن استخراج شده: {text.strip()}"
        else:
            return f"تحلیل تصویر: {description}"
    except Exception as e:
        return f"خطا در پردازش تصویر: {e}"

def analyze_text_file(file_obj, ext):
    ext = ext.lower()
    if ext == '.pdf':
        try:
            reader = PyPDF2.PdfReader(file_obj)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return f"متن استخراج شده از PDF: {text.strip()}"
        except Exception as e:
            return f"خطا در پردازش PDF: {e}"
    elif ext in ['.doc', '.docx']:
        try:
            doc = docx.Document(file_obj)
            text = "\n".join([para.text for para in doc.paragraphs])
            return f"متن استخراج شده از Word: {text.strip()}"
        except Exception as e:
            return f"خطا در پردازش فایل Word: {e}"
    else:
        try:
            content = file_obj.read().decode("utf-8")
            return f"متن استخراج شده: {content.strip()}"
        except Exception as e:
            return f"خطا در خواندن فایل متنی: {e}"

def analyze_audio(file_obj):
    try:
        import speech_recognition as sr
    except ImportError:
        return "امکان تحلیل صوتی فراهم نیست (کتابخانه speech_recognition نصب نشده است)."
    recognizer = sr.Recognizer()
    try:
        with sr.AudioFile(file_obj) as source:
            audio = recognizer.record(source)
        text = recognizer.recognize_google(audio, language="fa-IR")
        return f"متن استخراج شده از صدا: {text}"
    except Exception as e:
        return f"خطا در تشخیص صدا: {e}"

# ==================== توابع استفاده از مدل (Use_Model) ====================
def Use_Model(robot_name, mode_option):
    os.chdir('Models')
    folder_path = os.path.join(os.getcwd(), robot_name)
    if not os.path.exists(folder_path):
        print(f"Model {robot_name} Not Find!")
        exit()
    print(f"LOading {folder_path}...")
    model = GPT2LMHeadModel.from_pretrained(folder_path)
    tokenizer = GPT2Tokenizer.from_pretrained(folder_path)
    if tokenizer.pad_token is None:
        tokenizer.pad_token = tokenizer.eos_token
    device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
    model.to(device)
    history_file = os.path.join(folder_path, "chat_history.txt")
    if os.path.exists(history_file):
        with open(history_file, "r", encoding="utf-8") as f:
            chat_history = f.read()
    else:
        chat_history = ""

    def manage_history(chat_history, max_length=1024):
        if len(chat_history) > max_length:
            chat_history = chat_history[-max_length:]
        return chat_history

    def get_effective_max_length(input_ids, maxLength):
        max_context = model.config.n_ctx if hasattr(model.config, "n_ctx") else model.config.max_position_embeddings
        current_length = input_ids.shape[1]
        remaining = max_context - current_length
        effective_max = current_length + min(maxLength, remaining)
        return effective_max

    def generate_response(input_text, maxLength=10000, Seach_Mode=False, Argumentative_Mode=False,
                          Creative_Mode=False, DeepThink_Mode=False, Human_Mode=False, TextGenerate_Mode=False):
        if Seach_Mode:
            output = ""
            def get_top_links(query):
                return list(search(query, num=10))
            def extract_summary(url):
                try:
                    response = requests.get(url,timeout=4)
                    response.raise_for_status()
                    soup = BeautifulSoup(response.text, 'html.parser')
                    paragraphs = soup.find_all('p')
                    text = ' '.join([para.get_text() for para in paragraphs])
                    print("Site : ",url)
                    return text[:500]
                except Exception as e:
                    return f"Error fetching {url}: {e}"
            print("Start Search...")
            top_links = get_top_links(input_text)
            print("Finish Search...")
            for link in top_links:
                output += f"\nSite {link}\n Summary: {extract_summary(link)}\n"
            return output

        elif Argumentative_Mode:
            input_ids = tokenizer.encode(input_text, truncation=True, padding=True, max_length=1024, return_tensors="pt").to(device)
            attention_mask = input_ids.ne(tokenizer.pad_token_id).to(device)
            eff_max_length = get_effective_max_length(input_ids, maxLength)
            model.eval()
            with torch.no_grad():
                output = model.generate(
                    input_ids,
                    attention_mask=attention_mask,
                    max_length=eff_max_length,
                    num_return_sequences=1,
                    no_repeat_ngram_size=2,
                    top_p=0.0,
                    top_k=1,
                    temperature=0.1,
                    do_sample=False,
                    pad_token_id=tokenizer.eos_token_id,
                    eos_token_id=tokenizer.eos_token_id,
                )
            return tokenizer.decode(output[0], skip_special_tokens=True)

        elif Creative_Mode:
            input_ids = tokenizer.encode(input_text, truncation=True, padding=True, max_length=1024, return_tensors="pt").to(device)
            attention_mask = input_ids.ne(tokenizer.pad_token_id).to(device)
            eff_max_length = get_effective_max_length(input_ids, maxLength)
            model.eval()
            with torch.no_grad():
                output = model.generate(
                    input_ids,
                    attention_mask=attention_mask,
                    max_length=eff_max_length,
                    num_return_sequences=3,
                    no_repeat_ngram_size=3,
                    top_p=1.0,
                    top_k=50,
                    temperature=1.2,
                    do_sample=True,
                    pad_token_id=tokenizer.eos_token_id,
                )
            return tokenizer.decode(output[0], skip_special_tokens=True)

        elif DeepThink_Mode:
            input_text = f"Explore all the different possibilities and topics in different perspectives about this text: {input_text}"
            input_ids = tokenizer.encode(input_text, truncation=True, padding=True, max_length=1024, return_tensors="pt").to(device)
            attention_mask = input_ids.ne(tokenizer.pad_token_id).to(device)
            for token_id in input_ids[0]:
                if token_id >= tokenizer.vocab_size:
                    print(f"Token ID {token_id} is out of range.")
            print('Y')
            model.eval()
            with torch.no_grad():
                max_context = model.config.n_ctx if hasattr(model.config, "n_ctx") else model.config.max_position_embeddings
                current_length = input_ids.shape[1]
                remaining_tokens = max_context - current_length
                eff_new_tokens = min(200, remaining_tokens)
                output = model.generate(
                    input_ids,
                    attention_mask=attention_mask,
                    max_new_tokens=eff_new_tokens,
                    num_return_sequences=3,
                    no_repeat_ngram_size=2,
                    top_p=0.8,
                    top_k=20,
                    temperature=1.5,
                    do_sample=True,
                    pad_token_id=tokenizer.eos_token_id,
                )
            think_response = tokenizer.decode(output[0], skip_special_tokens=True)
            print('Y2')
            output_str = f"<think>\n{think_response}\n</think>"
            model.eval()
            with torch.no_grad():
                current_length = input_ids.shape[1]
                remaining_tokens = max_context - current_length
                eff_new_tokens = min(200, remaining_tokens)
                output_main = model.generate(
                    input_ids,
                    attention_mask=attention_mask,
                    max_new_tokens=eff_new_tokens,
                    num_return_sequences=1,
                    no_repeat_ngram_size=4,
                    top_p=0.9,
                    top_k=50,
                    temperature=0.85,
                    do_sample=True,
                    pad_token_id=tokenizer.eos_token_id,
                )
            responsemain = tokenizer.decode(output_main[0], skip_special_tokens=True)
            return f"{output_str}\n{responsemain}\n"

        elif Human_Mode:
            input_ids = tokenizer.encode(input_text, truncation=True, padding=True, max_length=1024, return_tensors="pt").to(device)
            attention_mask = input_ids.ne(tokenizer.pad_token_id).to(device)
            eff_max_length = get_effective_max_length(input_ids, maxLength)
            model.eval()
            with torch.no_grad():
                output = model.generate(
                    input_ids,
                    attention_mask=attention_mask,
                    max_length=eff_max_length,
                    num_return_sequences=1,
                    no_repeat_ngram_size=3,
                    top_p=0.95,
                    top_k=40,
                    temperature=0.9,
                    do_sample=True,
                    pad_token_id=tokenizer.eos_token_id,
                )
            return tokenizer.decode(output[0], skip_special_tokens=True)

        elif TextGenerate_Mode:
            input_ids = tokenizer.encode(input_text, truncation=True, padding=True, max_length=1024, return_tensors="pt").to(device)
            attention_mask = input_ids.ne(tokenizer.pad_token_id).to(device)
            model.eval()
            with torch.no_grad():
                eff_max_length = get_effective_max_length(input_ids, maxLength)
                output = model.generate(
                    input_ids,
                    attention_mask=attention_mask,
                    max_length=eff_max_length,
                    num_return_sequences=5,
                    no_repeat_ngram_size=4,
                    top_p=8,
                    top_k=50,
                    temperature=5,
                    do_sample=True,
                    pad_token_id=tokenizer.eos_token_id
                )
            return tokenizer.decode(output[0], skip_special_tokens=True)

        else:
            input_ids = tokenizer.encode(input_text, truncation=True, padding=True, max_length=1024, return_tensors="pt").to(device)
            attention_mask = input_ids.ne(tokenizer.pad_token_id).to(device)
            eff_max_length = get_effective_max_length(input_ids, maxLength)
            model.eval()
            with torch.no_grad():
                output = model.generate(
                    input_ids,
                    attention_mask=attention_mask,
                    max_length=eff_max_length,
                    num_return_sequences=1,
                    no_repeat_ngram_size=2,
                    top_p=0.7,
                    top_k=50,
                    temperature=50.7,
                    do_sample=True,
                    pad_token_id=tokenizer.eos_token_id
                )
            return tokenizer.decode(output[0], skip_special_tokens=True)

    # ==================== حالت چت ترمینالی ====================
    def runchat(modelName):
        his = str(input("Do you want use history? [Y]es/[N]o: "))
        use_history_flag = True if 'y' in his.lower() else False
        mode_input = str(input("[argumentative, creative, deepthink, human, textgenerate, search]\nEnter mode name if desired: ")).lower()
        while True:
            if use_history_flag:
                try:
                    with open(history_file, "r", encoding="utf-8") as f:
                        chat_history = f.read()
                except:
                    chat_history = ""
            flags = {}
            if mode_input == "argumentative":
                flags = {"Argumentative_Mode": True}
            elif mode_input == "creative":
                flags = {"Creative_Mode": True}
            elif mode_input == "deepthink":
                flags = {"DeepThink_Mode": True}
            elif mode_input == "human":
                flags = {"Human_Mode": True}
            elif mode_input == "textgenerate":
                flags = {"TextGenerate_Mode": True}
            elif mode_input == "search":
                flags = {"Seach_Mode": True}
            user_input = input("Your: ")
            if user_input.lower() == "/bye":
                print("Exit from chat...")
                break
            if use_history_flag:
                chat_history += f"you: {user_input}\n"
                chat_history = manage_history(chat_history)
                bot_response = generate_response(chat_history, **flags)
            else:
                bot_response = generate_response(user_input, **flags)
            if use_history_flag:
                chat_history += f"robot: {bot_response}\n"
            print(f"model {modelName}: {bot_response}")
            if use_history_flag:
                with open(history_file, "w", encoding="utf-8") as f:
                    f.write(chat_history)

    # ==================== حالت وب با Flask ====================
    def runweb():
        from flask import Flask, render_template, request, jsonify
        app = Flask(__name__)

        @app.route("/")
        def index():
            model_name = request.args.get("model", robot_name)
            # لیست مدل‌ها؛ به عنوان نمونه از پوشه Models استفاده می‌شود
            models = os.listdir(os.getcwd())
            return render_template("index.html", model_name=model_name, models=models)

        @app.route("/send", methods=["POST"])
        def send():
            mode = request.form.get("mode", "default").lower()
            user_message = request.form.get("message", "").strip()
            use_history = request.form.get("use_history", "true").lower() == "true"
            file = request.files.get("file")
            file_analysis = ""
            if file:
                filename = file.filename
                ext = os.path.splitext(filename)[1].lower()
                if ext in ['.png', '.jpg', '.jpeg', '.bmp', '.gif']:
                    file_analysis = analyze_image(file)
                elif ext in ['.pdf', '.doc', '.docx', '.xlsx', '.csv', '.txt']:
                    file_analysis = analyze_text_file(file, ext)
                elif ext in ['.mp3', '.wav', '.aac', '.ogg']:
                    file_analysis = analyze_audio(file)
                else:
                    file_analysis = f"فایل با پسوند {ext} پشتیبانی نمی‌شود."
                user_message += "\nتحلیل فایل: " + file_analysis
            if not user_message:
                return jsonify({"error": "Empty message"}), 400
            if use_history:
                try:
                    with open(history_file, "r", encoding="utf-8") as f:
                        current_history = f.read()
                except:
                    current_history = ""
                current_history += f"you: {user_message}\n"
            else:
                current_history = f"you: {user_message}\n"
            current_history = manage_history(current_history)
            flags = {}
            if mode == "argumentative":
                flags = {"Argumentative_Mode": True}
            elif mode == "creative":
                flags = {"Creative_Mode": True}
            elif mode == "deepthink":
                flags = {"DeepThink_Mode": True}
            elif mode == "human":
                flags = {"Human_Mode": True}
            elif mode == "textgenerate":
                flags = {"TextGenerate_Mode": True}
            elif mode == "search":
                flags = {"Seach_Mode": True}
            bot_response = generate_response(current_history, **flags)
            current_history += f"robot: {bot_response}\n"
            if use_history:
                with open(history_file, "w", encoding="utf-8") as f:
                    f.write(current_history)
            return jsonify({"response": bot_response})

        @app.route("/train", methods=["GET", "POST"])
        def train():
            from flask import redirect
            if request.method == "POST":
                robot_name_form = request.form.get("robot_name", "").strip()
                accuracy = float(request.form.get("accuracy", 0.9))
                test_size = float(request.form.get("test_size", 0.1))
                number_test = int(request.form.get("number_test", 3))
                files = request.files.getlist("files")
                if not robot_name_form or not files:
                    return "نام ربات و حداقل یک فایل انتخاب شود.", 400
                temp_dir = tempfile.mkdtemp()
                for f in files:
                    f.save(os.path.join(temp_dir, f.filename))
                try:
                    Create_Model(robot_name=robot_name_form, folder_path=temp_dir,
                                 accuracy_param=accuracy, test_size_param=test_size,
                                 number_test=number_test)
                    msg = f"مدل {robot_name_form} با موفقیت آموزش داده شد."
                except Exception as e:
                    msg = f"خطا در آموزش مدل: {e}"
                finally:
                    shutil.rmtree(temp_dir)
                return msg
            else:
                return render_template("train.html")
        app.run(host="0.0.0.0", port=5000, debug=False)

    # انتخاب حالت استفاده
    if mode_option == 'chat':
        runchat(robot_name)
    elif mode_option == 'web':
        runweb()

# ==================== فراخوانی برنامه اصلی ====================
try:
    p = sys.argv[1].lower()
    if p == 'list':
        try:
            os.chdir('Models')
            for n in os.listdir():
                print(" Model Name : ", n)
            os.chdir('..')
        except:
            os.mkdir('Models')
            print("You do not have any model! Create one!")
    elif p == 'runchat':
        Use_Model(robot_name=sys.argv[2], mode_option='chat')
    elif p == 'runweb':
        Use_Model(robot_name=sys.argv[2], mode_option='web')
    elif p == 'learnchat':
        # آموزش مدل در ترمینال
        try:
            sys.argv[6]
            Create_Model(robot_name=sys.argv[2], folder_path=sys.argv[3],
                         accuracy_param=sys.argv[4], number_test=sys.argv[5], test_size_param=sys.argv[6])
        except:
            Create_Model(robot_name=sys.argv[2], folder_path=sys.argv[3])
    elif p == 'learnweb':
        # آموزش مدل در وب
        from flask import Flask, render_template, request, redirect, url_for
        app = Flask(__name__)
        @app.route("/")
        def index():
            return render_template("train.html")
        @app.route("/train", methods=["POST"])
        def train():
            robot_name_form = request.form.get("robot_name", "").strip()
            accuracy = float(request.form.get("accuracy", 0.9))
            test_size = float(request.form.get("test_size", 0.1))
            number_test = int(request.form.get("number_test", 3))
            files = request.files.getlist("files")
            if not robot_name_form or not files:
                return "نام ربات و حداقل یک فایل انتخاب شود.", 400
            temp_dir = tempfile.mkdtemp()
            for f in files:
                f.save(os.path.join(temp_dir, f.filename))
            try:
                Create_Model(robot_name=robot_name_form, folder_path=temp_dir,
                             accuracy_param=accuracy, test_size_param=test_size,
                             number_test=number_test)
                msg = f"مدل {robot_name_form} با موفقیت آموزش داده شد."
            except Exception as e:
                msg = f"خطا در آموزش مدل: {e}"
            finally:
                shutil.rmtree(temp_dir)
            return msg
        app.run(host="0.0.0.0", port=5000, debug=False)
    else:
        raise Exception("Invalid command!")
except Exception as e:
    print(f"""
      
       P E R S I O N   A I
       
Coded by Mohhamad Taha Gorji
      
Create Models:
  persianai learnweb --> Crete new model (In Browser)
          
  persianai learnchat MODELNAME INFORMATION_FOLDER_PATH --> crete model (In Terminal)
          OR
    persianai learnchat MODELNAME INFORMATION_FOLDER_PATH accuracy_param test_size_param number_test --> For More Control
          
          
Use Models:
    persianai list --> show your models 
    persianai runweb MODELNAME --> run your model chat in browser
    persianai runchat MODEL_NAME --> run your model in terminal

Error: {e}
""")
